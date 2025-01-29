from pathlib import Path  # Для роботи з файловою системою
from docx import Document  # Бібліотека для створення та редагування Word-документів
import logging  # Ведення логування
import time  # Додавання затримок
import csv  # Читання CSV-файлів
import smtplib  # Відправка електронних листів
import ssl  # Захищене з'єднання для SMTP
from email.mime.base import MIMEBase  # Додавання вкладень у листи
from email.mime.multipart import MIMEMultipart  # Формування багатокомпонентного листа
from email import encoders  # Кодування вкладень
from email.mime.text import MIMEText  # Додавання текстового вмісту у лист
from email.header import Header  # Підтримка кодування заголовків

# Налаштування логування для відстеження процесу виконання
logging.basicConfig(level=logging.DEBUG, format=' %(asctime)s - %(levelname)s - %(message)s')
logging.debug('Початок виконання програми')

# Файли CSV зі списком регіонів та електронних адрес
QUERY_LIST_FILE_EMAILS = Path(f'./your_emails.csv')  # Файл зі списком email-адрес
QUERY_LIST_FILE_REGIONS = Path(f'./your_regions_list.csv')  # Файл зі списком регіонів

# Текст, який буде міститися у тілі електронного листа
body = "Доброго дня,\nНадсилаю інформаційний запит згідно ЗУ Про доступ до публічної інформації.\nПрошу зареєструвати та повідомити вхідний номер мого запиту.\nДякую за розуміння та співпрацю!\nЗ повагою,\n..."

def get_keywords(query_file):
    """Зчитує дані з CSV-файлу та повертає список значень (назви регіонів або email-адреси)"""
    with open(query_file, 'r') as i_file:
        rows = csv.reader(i_file, delimiter=',')
        keywords = [row[0] for row in rows]  # Зчитуємо перший стовпець у список
        return keywords

def get_para_data(output_doc_name, paragraph):
    """Копіює вміст та форматування абзацу у новий документ"""
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)  # Додаємо текст
        output_run.bold = run.bold  # Зберігаємо жирний шрифт
        output_run.italic = run.italic  # Зберігаємо курсив
        output_run.underline = run.underline  # Зберігаємо підкреслення
        output_run.font.color.rgb = run.font.color.rgb  # Зберігаємо колір тексту
        output_run.style.name = run.style.name  # Зберігаємо стиль шрифту
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment  # Вирівнювання тексту

def sendEmail(sending_file, file_name, email):
    """Відправляє електронний лист із вкладеним файлом"""
    time.sleep(2)  # Невелика пауза перед відправленням
    account = 'test.nikcenter.org'  # Логін поштового сервера
    sender = 'test@nikcenter.org'  # Адреса відправника
    sender_name = 'nikcenter'  # Ім'я відправника
    password = 'XXXXX'  # Пароль (необхідно замінити на реальний)
    
    logging.debug(f'Вхід у поштовий сервер для відправки {email}')
    context = ssl.create_default_context()
    smtpObj = smtplib.SMTP_SSL('mail.nikcenter.org', 465, context=context)  # Використовуємо захищене SSL-з'єднання
    smtpObj.ehlo()
    smtpObj.login(account, password)  # Вхід у поштовий обліковий запис
    
    msg = MIMEMultipart()
    msg['Subject'] = Header('Запит на доступ до публічної інформації', 'utf-8')  # Заголовок листа
    msg.attach(MIMEText(body.encode('utf-8'), _charset='utf-8'))  # Додаємо тіло листа
    
    with open(sending_file, 'rb') as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename= {file_name}")  # Додаємо вкладення
        msg.attach(part)
    
    logging.debug(f'Надсилання email до {email}')
    sendmailStatus = smtpObj.sendmail(sender, email, msg.as_string())
    
    if sendmailStatus != {}:
        logging.error(f'Помилка при відправці email до {email}: {sendmailStatus}')
    
    smtpObj.quit()

def formingDocx(i):
    """Створює Word-документ для конкретного регіону та надсилає його"""
    input_doc = Document('./zapyt.docx')  # Використовуємо шаблонний документ
    output_doc = Document()
    stringToPut = f'До ГУНП в {REGIONS[i]} області'  # Додаємо назву регіону
    paragraph = output_doc.add_paragraph(stringToPut)
    paragraph.alignment = 2  # Вирівнюємо текст справа
    
    for para in input_doc.paragraphs:
        get_para_data(output_doc, para)  # Копіюємо вміст шаблонного документа
    
    output_name = f'zapytGUNP{str(i)}.docx'  # Формуємо назву файлу
    output_name_fullPath = f'./{output_name}'  # Формуємо повний шлях
    output_doc.save(output_name_fullPath)  # Зберігаємо документ
    logging.debug(f"Документ збережено: {output_name}")
    
    sendEmail(output_name_fullPath, output_name, EMAILS[i])  # Надсилаємо документ

# Завантаження списків регіонів та email-адрес
REGIONS = get_keywords(QUERY_LIST_FILE_REGIONS)
EMAILS = get_keywords(QUERY_LIST_FILE_EMAILS)

# Генерація документів та відправка листів
for i in range(len(EMAILS)):
    formingDocx(i)  # Створюємо документ і відправляємо його

logging.debug("Виконання завершено!")
print("Готово!")
